"""Извлечение текста из файлов обычными бэкенд-библиотеками (по постановке:
в модель отдаём уже текст). PDF -> pypdf, DOCX -> python-docx, txt/md/html -> как есть.

Старый бинарный .doc (Word 97-2003) НЕ поддерживается: python-docx его не читает,
а попытка прочитать как текст даёт мусор на выходе фильтра — что опасно
(мусор "успешно чистится", реальное ПД может утечь). Поэтому — явная ошибка."""
import io
import re

DOC_SIGNATURE = b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"  # OLE2: Word 97-2003 и прочий старый Office


def extract_text(filename: str, content: bytes) -> str:
    name = filename.lower()
    if name.endswith(".doc") or content[:8] == DOC_SIGNATURE:
        raise ValueError(
            "Формат .doc (Word 97-2003) не поддерживается — пересохраните файл в .docx"
        )
    if name.endswith(".pdf"):
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(content))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    if name.endswith(".docx"):
        import docx
        d = docx.Document(io.BytesIO(content))
        parts = [p.text for p in d.paragraphs]
        for table in d.tables:
            for row in table.rows:
                parts.append(" | ".join(c.text for c in row.cells))
        return "\n".join(parts)
    if name.endswith((".html", ".htm")):
        text = content.decode("utf-8", errors="replace")
        text = re.sub(r"<script.*?</script>|<style.*?</style>", " ", text, flags=re.DOTALL | re.IGNORECASE)
        return re.sub(r"<[^>]+>", " ", text)
    return content.decode("utf-8", errors="replace")  # txt, md и всё остальное
